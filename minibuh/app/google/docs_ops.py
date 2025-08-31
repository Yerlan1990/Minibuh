from docx import Document
from io import BytesIO
from typing import Dict

def render_docx_from_context(title: str, context: Dict[str, str]) -> bytes:
    doc = Document()
    doc.add_heading(title, 0)
    for k, v in context.items():
        doc.add_paragraph(f"{k}: {v}")
    b = BytesIO()
    doc.save(b)
    return b.getvalue()
